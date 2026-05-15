# this file is fully vibecoded
from docx import Document
from collections import defaultdict
import datetime
import io
import re

class tableGen(object):
	"""Класс для генерации таблиц и кастомной обработки в Word"""
	def __init__(self, arg=None):
		super(tableGen, self).__init__()
		self.arg = arg

	def _parse_date(self, date_str):
		try:
			return datetime.datetime.strptime(str(date_str).strip(), "%d.%m.%Y")
		except Exception:
			return None

	def make_grouped_data(self, data_list):
		groups = defaultdict(list)
		for item in data_list:
			uk = str(item.get("УК", "Неизвестная УК"))
			groups[uk].append(item)
		
		result = []
		for uk, items in groups.items():
			dates = []
			for i in items:
				d_str = i.get("Дата", "")
				d_obj = self._parse_date(d_str)
				if d_obj:
					dates.append(d_obj)
			
			if dates:
				min_d = min(dates)
				max_d = max(dates)
				if min_d == max_d:
					date_label = min_d.strftime("%d.%m.%Y")
				else:
					date_label = f"с {min_d.strftime('%d.%m.%Y')} по {max_d.strftime('%d.%m.%Y')}"
			else:
				date_label = str(items[0].get("Дата", ""))

			base_item = items[0].copy()
			base_item["Дата"] = date_label
			base_item["__table_records__"] = items
			result.append(base_item)
			
		return result

	def castom_transform_addres(self, item, config_path=None):
		"""
		Пользовательская функция для сборки и нормализации адреса.
		Вы можете изменить эту логику позже.
		"""
		try:
			from .load import _transform_address
		except ImportError:
			from load import _transform_address

		def is_valid(val):
			if val is None:
				return False
			s = str(val).strip()
			return bool(s and s.lower() not in ['nan', 'none'])

		addr_obj_raw = str(item.get("Адрес объекта", item.get("Адрес", item.get("Адрес_объекта", "")))).strip()
		
		# Делаем каждое слово заглавным, но сохраняем регистр для маркеров
		words = addr_obj_raw.split()
		capitalized_words = []
		for w in words:
			if w.lower() in ["д.", "кв.", "ул.", "г.", "с.", "п.", "рп."]:
				capitalized_words.append(w.lower())
			else:
				# Обработка дефисов: РОСТОВ-НА-ДОНУ -> Ростов-На-Дону
				if "-" in w:
					capitalized_words.append("-".join(part.capitalize() for part in w.split("-")))
				else:
					capitalized_words.append(w.capitalize())
		
		addr_obj = " ".join(capitalized_words)
		
		# Проверяем, является ли адрес "полным" (содержит запятые или маркеры д. / кв.)
		is_full = "," in addr_obj or " д." in addr_obj.lower() or " кв." in addr_obj.lower()
		
		if is_full:
			# Если это уже готовая строка, не добавляем "г. НП" и "ул." вручную
			addr_str = addr_obj
			# Но если дом и квартира лежат в отдельных полях, докидываем их
			if is_valid(item.get("Дом")) and f"д. {item['Дом']}" not in addr_str:
				addr_str += f", д. {item['Дом']}"
			if is_valid(item.get("Квартира")) and f"кв. {item['Квартира']}" not in addr_str:
				kv = str(item['Квартира'])
				if kv.endswith('.0'): kv = kv[:-2]
				addr_str += f", кв. {kv}"
				
			addr_norm = _transform_address(addr_str, config_path=config_path)
			return addr_norm if addr_norm else addr_str

		# Если это просто название улицы (например "ВОСХОД УЛ." или "КИРОВА")
		import re
		clean_street = re.sub(r'(?i)\bул(?:ица)?\.?', '', addr_obj).strip()
		
		if clean_street:
			addr_str = f"ул. {clean_street}"
		else:
			addr_str = ""

		# Добавляем дом
		if is_valid(item.get("Дом")):
			дом = str(item['Дом'])
			if дом.endswith('.0'): дом = дом[:-2]
			addr_str += f", д. {дом}"
			
		# Добавляем квартиру
		if is_valid(item.get("Квартира")):
			кв = str(item['Квартира'])
			if кв.endswith('.0'): кв = кв[:-2]
			addr_str += f", кв. {кв}"
		
		# Обработка поля НП
		np_val = str(item.get("НП", "")).strip()
		np_str = ""
		if np_val and is_valid(np_val):
			np_words = np_val.split()
			np_clean_words = []
			prefix = "г." # По умолчанию
			
			for w in np_words:
				w_lower = w.lower().strip(".")
				if w_lower in ["с", "село"]:
					prefix = "с."
				elif w_lower in ["п", "пос", "поселок", "посёлок"]:
					prefix = "п."
				elif w_lower in ["рп", "рабпос"]:
					prefix = "рп."
				elif w_lower in ["г", "город"]:
					prefix = "г."
				else:
					np_clean_words.append(w.capitalize())
			
			if np_clean_words:
				np_str = f"{prefix} {' '.join(np_clean_words)}"

		if np_str:
			if addr_str:
				addr_str = f"{np_str}, {addr_str}"
			else:
				addr_str = np_str
		
		addr_norm = _transform_address(addr_str, config_path=config_path)
		return addr_norm if addr_norm else addr_str

	def process_table_logic(self, data_list, path_to_shablon, output_path, shablon_module, config_path=None):
		"""
		Точка входа из shablon.py. Читает документ, если там есть {{:table}},
		группирует данные и выполняет рендер.
		"""
		from docx.shared import Cm
		doc = Document(path_to_shablon)
		has_table = False
		for p in doc.paragraphs:
			if '{{:table}}' in p.text:
				has_table = True
				break
		
		if not has_table:
			return False

		grouped_data = self.make_grouped_data(data_list)
		
		from docxcompose.composer import Composer
		composer = None
		
		for idx, group_item in enumerate(grouped_data):
			doc = Document(path_to_shablon)
			
			target_p = None
			for p in doc.paragraphs:
				if '{{:table}}' in p.text:
					target_p = p
					break
			
			if target_p:
				target_p.text = target_p.text.replace('{{:table}}', '')
				
				records = group_item["__table_records__"]
				table = doc.add_table(rows=1, cols=3)
				table.style = 'Table Grid'
				table.autofit = False
				table.allow_autofit = False
				
				# Задаем ширину колонок
				widths = [Cm(1.5), Cm(4.5), Cm(10.5)]
				for j, col in enumerate(table.columns):
					col.width = widths[j]
				
				tbl_p = target_p._p
				tbl_p.addnext(table._tbl)

				hdr_cells = table.rows[0].cells
				hdr_cells[0].text = "№"
				hdr_cells[1].text = "Лицевой счет абонента"
				hdr_cells[2].text = "Полный адрес (город, улица, дом, квартира)"
				
				for j, cell in enumerate(hdr_cells):
					cell.width = widths[j]
				
				for i, item in enumerate(records):
					row_cells = table.add_row().cells
					for j, cell in enumerate(row_cells):
						cell.width = widths[j]

					row_cells[0].text = str(i + 1)
					
					# Поддержка ключа с пробелом и с подчеркиванием
					lc = item.get("Лицевой_счет", item.get("Лицевой счет", item.get("Лицевой счет абонента", "")))
					row_cells[1].text = str(lc)
					
					addr_norm = self.castom_transform_addres(item, config_path=config_path)
					
					date_val = str(item.get("Дата", ""))
					time_val = str(item.get("Время", ""))
					if not time_val:
						for k, v in item.items():
							if isinstance(k, str) and 'Unnamed' in k and re.match(r'^\d{2}:\d{2}-\d{2}:\d{2}$', str(v)):
								time_val = str(v)
								break
					
					datetime_str = f"{date_val}"
					if time_val:
						datetime_str += f" в период {time_val} часов"
						
					row_cells[2].text = f"{addr_norm}\n{datetime_str}"

			stream = io.BytesIO()
			doc.save(stream)
			stream.seek(0)
			
			tpl = shablon_module.MyDocxTemplate(stream)
			shablon_module._safe_render(tpl, group_item, config_path=config_path)
			
			out_stream = io.BytesIO()
			tpl.save(out_stream)
			out_stream.seek(0)
			
			if composer is None:
				composer = Composer(Document(out_stream))
			else:
				composer.doc.add_page_break()
				composer.append(Document(out_stream))
				
		import time as t
		from pathlib import Path
		output_path = Path(output_path)
		if output_path.exists():
			new_stem = f"{output_path.stem}_{int(t.time())}"
			final_path = output_path.with_stem(new_stem)
		else:
			final_path = output_path

		composer.save(final_path)
		print("Done grouped tables!")
		return final_path

def main():
	pass

if __name__ == '__main__':
	main()